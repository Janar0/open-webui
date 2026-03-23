import logging
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
from html import escape

from open_webui.models.chats import ChatTitleMessagesForm

log = logging.getLogger(__name__)

# Styling constants
FONT_BODY = "Calibri"
FONT_CODE = "Courier New"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_CODE = Pt(9)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_ROLE = Pt(12)
COLOR_USER = RGBColor(0x25, 0x63, 0xEB)  # Blue
COLOR_ASSISTANT = RGBColor(0x05, 0x96, 0x69)  # Green
COLOR_TIMESTAMP = RGBColor(0x6B, 0x72, 0x80)  # Gray
COLOR_CODE_BG = "F3F4F6"
COLOR_TABLE_HEADER = "3B82F6"
COLOR_BLOCKQUOTE_BG = "F9FAFB"
COLOR_BLOCKQUOTE_BORDER = "3B82F6"


class DocxGenerator:
    def __init__(self, form_data: ChatTitleMessagesForm):
        self.form_data = form_data

    def format_timestamp(self, timestamp: float) -> str:
        try:
            date_time = datetime.fromtimestamp(timestamp)
            return date_time.strftime("%Y-%m-%d, %H:%M:%S")
        except (ValueError, TypeError):
            return ""

    def generate_chat_docx(self) -> bytes:
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_BODY
        font.size = FONT_SIZE_BODY
        font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

        # Title
        title_para = doc.add_heading(self.form_data.title, level=1)
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        # Thin separator after title
        self._add_horizontal_rule(doc)

        # Messages
        for i, message in enumerate(self.form_data.messages):
            self._add_message(doc, message)
            if i < len(self.form_data.messages) - 1:
                self._add_horizontal_rule(doc)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_message_docx(self) -> bytes:
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_BODY
        font.size = FONT_SIZE_BODY
        font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

        if self.form_data.messages:
            self._add_message(doc, self.form_data.messages[0])

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _add_message(self, doc: Document, message: Dict[str, Any]):
        role = message.get("role", "user")
        content = message.get("content", "")
        timestamp = message.get("timestamp")
        model = message.get("model", "") if role == "assistant" else ""

        # Role label
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(role.title())
        role_run.bold = True
        role_run.font.size = FONT_SIZE_ROLE
        role_run.font.color.rgb = COLOR_USER if role == "user" else COLOR_ASSISTANT

        if model:
            model_run = role_para.add_run(f"  ({model})")
            model_run.font.size = FONT_SIZE_SMALL
            model_run.font.color.rgb = COLOR_TIMESTAMP

        # Timestamp
        if timestamp:
            date_str = self.format_timestamp(timestamp)
            if date_str:
                ts_para = doc.add_paragraph()
                ts_run = ts_para.add_run(date_str)
                ts_run.font.size = FONT_SIZE_SMALL
                ts_run.font.color.rgb = COLOR_TIMESTAMP
                ts_para.paragraph_format.space_after = Pt(2)

        # Images from message files
        files = message.get("files", [])
        if files:
            for file_info in files:
                if file_info.get("type") == "image" or (
                    file_info.get("content_type", "")
                ).startswith("image/"):
                    self._add_image(doc, file_info)

        # Parse and render markdown content
        if content:
            self._parse_markdown(doc, content)

    def _add_image(self, doc: Document, file_info: Dict[str, Any]):
        url = file_info.get("url", "")
        if not url:
            return

        try:
            # Handle relative URLs (internal files)
            if url.startswith("/"):
                # Internal file - try to read directly from storage
                self._add_image_placeholder(
                    doc, file_info.get("name", "image")
                )
                return

            # Handle base64 data URIs
            if url.startswith("data:image"):
                match = re.match(r"data:image/\w+;base64,(.+)", url)
                if match:
                    import base64

                    img_data = base64.b64decode(match.group(1))
                    img_stream = BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(5))
                    return

            # External URL - fetch
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                img_stream = BytesIO(response.content)
                doc.add_picture(img_stream, width=Inches(5))
                return
        except Exception as e:
            log.warning(f"Failed to embed image: {e}")

        self._add_image_placeholder(doc, file_info.get("name", "image"))

    def _add_image_placeholder(self, doc: Document, alt_text: str):
        para = doc.add_paragraph()
        run = para.add_run(f"[Image: {alt_text}]")
        run.italic = True
        run.font.color.rgb = COLOR_TIMESTAMP

    def _add_horizontal_rule(self, doc: Document):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D1D5DB")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _parse_markdown(self, doc: Document, content: str):
        # Split content into blocks: code blocks vs regular text
        parts = re.split(r"(```[\s\S]*?```)", content)

        for part in parts:
            if part.startswith("```") and part.endswith("```"):
                self._add_code_block(doc, part)
            else:
                self._parse_regular_text(doc, part)

    def _parse_regular_text(self, doc: Document, text: str):
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Heading
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                h = doc.add_heading(level=min(level, 4))
                self._add_formatted_runs(h, heading_text)
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^[-*_]{3,}\s*$", stripped):
                self._add_horizontal_rule(doc)
                i += 1
                continue

            # Table detection
            if "|" in stripped and i + 1 < len(lines):
                table_lines = []
                j = i
                while j < len(lines) and "|" in lines[j].strip():
                    table_lines.append(lines[j].strip())
                    j += 1
                if len(table_lines) >= 2 and re.match(
                    r"^\|?\s*[-:]+[-|:\s]+\s*\|?$", table_lines[1]
                ):
                    self._add_table(doc, table_lines)
                    i = j
                    continue

            # Blockquote
            if stripped.startswith(">"):
                quote_lines = []
                j = i
                while j < len(lines) and lines[j].strip().startswith(">"):
                    quote_lines.append(
                        re.sub(r"^>\s?", "", lines[j].strip())
                    )
                    j += 1
                self._add_blockquote(doc, "\n".join(quote_lines))
                i = j
                continue

            # Unordered list
            list_match = re.match(r"^[-*+]\s+(.+)$", stripped)
            if list_match:
                para = doc.add_paragraph(style="List Bullet")
                self._add_formatted_runs(para, list_match.group(1))
                i += 1
                continue

            # Ordered list
            ol_match = re.match(r"^\d+\.\s+(.+)$", stripped)
            if ol_match:
                para = doc.add_paragraph(style="List Number")
                self._add_formatted_runs(para, ol_match.group(1))
                i += 1
                continue

            # Markdown image
            img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
            if img_match:
                alt_text = img_match.group(1)
                img_url = img_match.group(2)
                self._add_image(doc, {"url": img_url, "name": alt_text})
                i += 1
                continue

            # Regular paragraph
            para = doc.add_paragraph()
            self._add_formatted_runs(para, stripped)
            i += 1

    def _add_formatted_runs(self, paragraph, text: str):
        # Pattern to match inline markdown: bold+italic, bold, italic, code, links, images
        pattern = re.compile(
            r"(\*\*\*(.+?)\*\*\*)"  # bold+italic
            r"|(\*\*(.+?)\*\*)"  # bold
            r"|(__(.+?)__)"  # bold (underscores)
            r"|(\*(.+?)\*)"  # italic
            r"|(_([^_]+?)_)"  # italic (underscores)
            r"|(`([^`]+?)`)"  # inline code
            r"|(\[([^\]]+?)\]\(([^)]+?)\))"  # link
            r"|(!\[([^\]]*?)\]\(([^)]+?)\))"  # image (inline)
        )

        last_end = 0
        for match in pattern.finditer(text):
            # Add plain text before this match
            if match.start() > last_end:
                plain = text[last_end : match.start()]
                if plain:
                    run = paragraph.add_run(plain)
                    run.font.name = FONT_BODY
                    run.font.size = FONT_SIZE_BODY

            if match.group(2):  # bold+italic
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY
            elif match.group(4):  # bold **
                run = paragraph.add_run(match.group(4))
                run.bold = True
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY
            elif match.group(6):  # bold __
                run = paragraph.add_run(match.group(6))
                run.bold = True
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY
            elif match.group(8):  # italic *
                run = paragraph.add_run(match.group(8))
                run.italic = True
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY
            elif match.group(10):  # italic _
                run = paragraph.add_run(match.group(10))
                run.italic = True
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY
            elif match.group(12):  # inline code
                run = paragraph.add_run(match.group(12))
                run.font.name = FONT_CODE
                run.font.size = FONT_SIZE_CODE
                self._set_run_shading(run, COLOR_CODE_BG)
            elif match.group(14):  # link [text](url)
                link_text = match.group(14)
                link_url = match.group(15)
                self._add_hyperlink(paragraph, link_url, link_text)
            elif match.group(17):  # inline image
                # Just add alt text for inline images
                alt = match.group(17) or "image"
                run = paragraph.add_run(f"[Image: {alt}]")
                run.italic = True
                run.font.color.rgb = COLOR_TIMESTAMP

            last_end = match.end()

        # Add remaining plain text
        if last_end < len(text):
            remaining = text[last_end:]
            if remaining:
                run = paragraph.add_run(remaining)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY

    def _set_run_shading(self, run, color: str):
        rPr = run._element.get_or_add_rPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        rPr.append(shading)

    def _add_hyperlink(self, paragraph, url: str, text: str):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        c = OxmlElement("w:color")
        c.set(qn("w:val"), "2563EB")
        rPr.append(c)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_BODY)
        rFonts.set(qn("w:hAnsi"), FONT_BODY)
        rPr.append(rFonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")  # 11pt = 22 half-points
        rPr.append(sz)

        new_run.append(rPr)
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._element.append(hyperlink)

    def _add_code_block(self, doc: Document, code_block: str):
        # Extract language and code
        lines = code_block.split("\n")
        first_line = lines[0].strip("`").strip()
        language = first_line if first_line else ""
        code_lines = lines[1:-1] if len(lines) > 2 else []
        code_text = "\n".join(code_lines)

        # Create a single-cell table for the code block background
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.cell(0, 0)

        # Set cell shading
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), COLOR_CODE_BG)
        tcPr.append(shading)

        # Set cell margins for padding
        tcMar = OxmlElement("w:tcMar")
        for side in ["top", "left", "bottom", "right"]:
            margin = OxmlElement(f"w:{side}")
            margin.set(qn("w:w"), "120")
            margin.set(qn("w:type"), "dxa")
            tcMar.append(margin)
        tcPr.append(tcMar)

        # Language label
        if language:
            lang_para = cell.paragraphs[0]
            lang_run = lang_para.add_run(language)
            lang_run.font.name = FONT_CODE
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = COLOR_TIMESTAMP
            lang_run.bold = True
            lang_para.paragraph_format.space_after = Pt(4)

            # Code content in new paragraph
            code_para = cell.add_paragraph()
        else:
            code_para = cell.paragraphs[0]

        code_run = code_para.add_run(code_text)
        code_run.font.name = FONT_CODE
        code_run.font.size = FONT_SIZE_CODE
        code_run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        code_para.paragraph_format.space_after = Pt(0)
        code_para.paragraph_format.line_spacing = 1.0

        # Set table borders to thin gray
        tbl = table._element
        tblPr = tbl.get_or_add_tblPr()
        borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "D1D5DB")
            borders.append(border)
        tblPr.append(borders)

        # Make table full width
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")  # 100% width
        tblPr.append(tblW)

        # Add small spacing after code block
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(4)

    def _add_table(self, doc: Document, table_lines: List[str]):
        # Parse header
        header_cells = [
            c.strip() for c in table_lines[0].strip("|").split("|")
        ]
        # Skip separator line (index 1)
        data_rows = []
        for line in table_lines[2:]:
            cells = [c.strip() for c in line.strip("|").split("|")]
            data_rows.append(cells)

        num_cols = len(header_cells)
        num_rows = len(data_rows) + 1  # +1 for header

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, header_text in enumerate(header_cells):
            if j < num_cols:
                cell = table.cell(0, j)
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(header_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY

                # Blue background for header
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), COLOR_TABLE_HEADER)
                tcPr.append(shading)

        # Data rows
        for i, row_data in enumerate(data_rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.cell(i + 1, j)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    self._add_formatted_runs(para, cell_text)

        # Add spacing after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(4)

    def _add_blockquote(self, doc: Document, text: str):
        # Use a single-cell table for blockquote styling
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)

        # Cell shading
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), COLOR_BLOCKQUOTE_BG)
        tcPr.append(shading)

        # Left border (blue accent)
        tcBorders = OxmlElement("w:tcBorders")
        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "24")  # Thick left border
        left_border.set(qn("w:space"), "0")
        left_border.set(qn("w:color"), COLOR_BLOCKQUOTE_BORDER)
        tcBorders.append(left_border)
        # Thin or no borders for other sides
        for side in ["top", "right", "bottom"]:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            tcBorders.append(border)
        tcPr.append(tcBorders)

        # Cell padding
        tcMar = OxmlElement("w:tcMar")
        for side_name in ["top", "left", "bottom", "right"]:
            margin = OxmlElement(f"w:{side_name}")
            margin.set(qn("w:w"), "120")
            margin.set(qn("w:type"), "dxa")
            tcMar.append(margin)
        tcPr.append(tcMar)

        # Content
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.italic = True
        run.font.name = FONT_BODY
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

        # Full width table, no outer borders
        tbl = table._element
        tblPr = tbl.get_or_add_tblPr()

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")
        tblPr.append(tblW)

        borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            borders.append(border)
        tblPr.append(borders)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(4)
